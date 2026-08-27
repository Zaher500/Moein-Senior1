import io
import os
import re
import arabic_reshaper
from bidi.algorithm import get_display
import unicodedata

# =========================================================
# PDF imports
# =========================================================

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate,
    Spacer,
    Flowable,
)

# =========================================================
# Word imports
# =========================================================

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================================================
# TEXT / DIRECTION HELPERS
# =========================================================

ARABIC_PATTERN = re.compile(r"[\u0600-\u06FF]")
# LRM = "\u200E"   # Left-to-Right Mark
# RLM = "\u200F"   # Right-to-Left Mark
# LRI = "\u2066"   # Left-to-Right Isolate
# RLI = "\u2067"   # Right-to-Left Isolate
# PDI = "\u2069"   # Pop Directional Isolate
# LTR_NEUTRAL_CHARS = set("()[]{}<>:/\\-–—_.;,+=#%&*@!?")

def contains_arabic(text: str) -> bool:
    """
    Check whether text contains Arabic characters.
    """
    return bool(ARABIC_PATTERN.search(text))


def detect_base_direction(text: str) -> str:
    """
    Detect paragraph direction from the first strong letter.

    Returns:
        R -> Arabic / RTL
        L -> English / LTR
    """

    for char in text:

        if re.match(r"[\u0600-\u06FF]", char):
            return "R"

        if char.isalpha():
            return "L"

    return "L"


def remove_markdown(text: str) -> str:
    """
    Remove simple Markdown markers for PDF.

    Example:
        **CMMI** -> CMMI
    """

    text = re.sub(
        r"\*\*(.*?)\*\*",
        r"\1",
        text
    )

    text = re.sub(
        r"__(.*?)__",
        r"\1",
        text
    )

    return text


# =========================================================
# PDF FONT
# =========================================================

def register_pdf_font():
    """
    Register a Unicode font that supports Arabic.

    Tries common Windows and Linux fonts.
    """

    possible_fonts = [

        # Windows
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",

        # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]

    for font_path in possible_fonts:

        if os.path.exists(font_path):

            pdfmetrics.registerFont(
                TTFont(
                    "SummaryFont",
                    font_path
                )
            )

            return "SummaryFont"

    raise FileNotFoundError(
        "No Arabic-compatible font was found. "
        "Please install Arial, Tahoma, "
        "Segoe UI, DejaVu Sans, or FreeSans."
    )


# =========================================================
# PDF TEXT PREPARATION
# =========================================================

def prepare_visual_text(
    text: str,
    base_direction: str
) -> str:
    """
    Convert ONE logical line into visual PDF text.

    Arabic reshaping + BiDi are performed after
    the text has already been wrapped into a line.

    This prevents Arabic paragraph lines from
    appearing in reverse vertical order.
    """

    if not text:
        return ""

    # Pure English / LTR
    if not contains_arabic(text):
        return text

    # Connect Arabic characters
    reshaped = arabic_reshaper.reshape(text)

    # Handle Arabic + English mixed direction
    visual = get_display(
        reshaped,
        base_dir=base_direction
    )

    return visual


def visual_width(
    text: str,
    font_name: str,
    font_size: float,
    base_direction: str
) -> float:
    """
    Calculate the real width of the text
    after Arabic shaping and BiDi processing.
    """

    visual = prepare_visual_text(
        text,
        base_direction
    )

    return pdfmetrics.stringWidth(
        visual,
        font_name,
        font_size
    )


# =========================================================
# PDF TEXT WRAPPING
# =========================================================

def wrap_logical_text(
    text: str,
    max_width: float,
    font_name: str,
    font_size: float,
    base_direction: str
):
    """
    Wrap text BEFORE applying final BiDi rendering.

    Important:
    Each resulting line is later processed separately.

    This prevents the problem where a complete Arabic
    paragraph is reversed first and then wrapped by ReportLab.
    """

    words = text.split()

    if not words:
        return [""]

    lines = []

    current_line = words[0]

    for word in words[1:]:

        candidate = f"{current_line} {word}"

        candidate_width = visual_width(
            candidate,
            font_name,
            font_size,
            base_direction
        )

        if candidate_width <= max_width:

            current_line = candidate

        else:

            lines.append(current_line)

            current_line = word

    if current_line:

        lines.append(current_line)

    return lines


# =========================================================
# CUSTOM PDF TEXT BLOCK
# =========================================================

class BidiTextBlock(Flowable):
    """
    Custom ReportLab Flowable for:

        - Arabic
        - English
        - Mixed Arabic + English

    The important difference from Paragraph:

    Paragraph may wrap the text again after we process BiDi.

    This class wraps exactly once, then draws every line
    directly on the PDF canvas.
    """

    def __init__(
        self,
        text,
        font_name,
        font_size=11,
        leading=17,
        space_after=6,
        direction=None
    ):

        super().__init__()

        self.text = text

        self.font_name = font_name

        self.font_size = font_size

        self.leading = leading

        self.space_after = space_after

        self.direction = (
            direction
            or detect_base_direction(text)
        )

        self.lines = []

        self.available_width = 0

    def wrap(
        self,
        avail_width,
        avail_height
    ):

        self.available_width = avail_width

        self.lines = wrap_logical_text(
            text=self.text,
            max_width=avail_width,
            font_name=self.font_name,
            font_size=self.font_size,
            base_direction=self.direction
        )

        self.width = avail_width

        self.height = (
            len(self.lines) * self.leading
            + self.space_after
        )

        return (
            self.width,
            self.height
        )

    def draw(self):

        canvas = self.canv

        canvas.setFont(
            self.font_name,
            self.font_size
        )

        y = (
            self.height
            - self.leading
        )

        for logical_line in self.lines:

            visual_line = prepare_visual_text(
                logical_line,
                self.direction
            )

            # =============================================
            # Arabic / RTL
            # =============================================

            if self.direction == "R":

                canvas.drawRightString(
                    self.available_width,
                    y,
                    visual_line
                )

            # =============================================
            # English / LTR
            # =============================================

            else:

                canvas.drawString(
                    0,
                    y,
                    visual_line
                )

            y -= self.leading


# =========================================================
# CUSTOM PDF LIST ITEM
# =========================================================

class BidiListItem(Flowable):
    """
    Custom bullet / numbered list item.

    The marker (• / 1. / 2.) is drawn separately
    from the actual text.

    This prevents BiDi from moving the marker to
    the wrong side.
    """

    def __init__(
        self,
        text,
        marker,
        font_name,
        font_size=11,
        leading=17,
        space_after=5
    ):

        super().__init__()

        self.text = text

        self.marker = marker

        self.font_name = font_name

        self.font_size = font_size

        self.leading = leading

        self.space_after = space_after

        self.direction = detect_base_direction(
            text
        )

        self.lines = []

    def wrap(
        self,
        avail_width,
        avail_height
    ):

        self.available_width = avail_width

        # Width of:
        # •
        # 1.
        # 2.
        self.marker_width = pdfmetrics.stringWidth(
            self.marker,
            self.font_name,
            self.font_size
        )

        # Space between marker and text
        self.gap = 10

        self.text_width = (
            avail_width
            - self.marker_width
            - self.gap
        )

        self.lines = wrap_logical_text(
            text=self.text,
            max_width=self.text_width,
            font_name=self.font_name,
            font_size=self.font_size,
            base_direction=self.direction
        )

        self.width = avail_width

        self.height = (
            len(self.lines) * self.leading
            + self.space_after
        )

        return (
            self.width,
            self.height
        )

    def draw(self):

        canvas = self.canv

        canvas.setFont(
            self.font_name,
            self.font_size
        )

        y = (
            self.height
            - self.leading
        )

        # =================================================
        # Arabic list
        # =================================================

        if self.direction == "R":

            # Marker goes to far right
            canvas.drawRightString(
                self.available_width,
                y,
                self.marker
            )

            text_right_edge = (
                self.available_width
                - self.marker_width
                - self.gap
            )

            for logical_line in self.lines:

                visual_line = prepare_visual_text(
                    logical_line,
                    "R"
                )

                canvas.drawRightString(
                    text_right_edge,
                    y,
                    visual_line
                )

                y -= self.leading

        # =================================================
        # English list
        # =================================================

        else:

            # Marker goes to left
            canvas.drawString(
                0,
                y,
                self.marker
            )

            text_left_edge = (
                self.marker_width
                + self.gap
            )

            for logical_line in self.lines:

                visual_line = prepare_visual_text(
                    logical_line,
                    "L"
                )

                canvas.drawString(
                    text_left_edge,
                    y,
                    visual_line
                )

                y -= self.leading


# =========================================================
# PDF EXPORT
# =========================================================

def generate_pdf(summary_text: str):
    """
    Generate PDF from stored summary text.

    Supports:
        - Arabic
        - English
        - Arabic + English mixed text
        - Headings
        - Bullet lists
        - Numbered lists
    """

    buffer = io.BytesIO()

    font_name = register_pdf_font()

    document = SimpleDocTemplate(
        buffer,

        pagesize=A4,

        rightMargin=50,
        leftMargin=50,

        topMargin=50,
        bottomMargin=50,

        title="Lecture Summary"
    )

    elements = []

    # =====================================================
    # Process summary line-by-line
    # =====================================================

    for raw_line in summary_text.splitlines():

        line = raw_line.strip()

        # =================================================
        # Empty line
        # =================================================

        if not line:

            elements.append(
                Spacer(
                    1,
                    8
                )
            )

            continue

        # =================================================
        # Detect Markdown heading
        # =================================================

        heading_level = None

        if line.startswith("### "):

            heading_level = 3

            line = line[4:]

        elif line.startswith("## "):

            heading_level = 2

            line = line[3:]

        elif line.startswith("# "):

            heading_level = 1

            line = line[2:]

        # =================================================
        # Detect bullet
        # =================================================

        is_bullet = False

        if line.startswith(
            (
                "- ",
                "* ",
                "• "
            )
        ):

            is_bullet = True

            line = line[2:].strip()

        # =================================================
        # Detect numbered list
        # =================================================

        numbered_match = re.match(
            r"^(\d+)\.\s+(.*)",
            line
        )

        number = None

        if numbered_match:

            number = numbered_match.group(1)

            line = numbered_match.group(2)

        # =================================================
        # Remove **markdown**
        # =================================================

        line = remove_markdown(
            line
        )

        # =================================================
        # Heading
        # =================================================

        if heading_level:

            if heading_level == 1:

                font_size = 18

                leading = 24

            elif heading_level == 2:

                font_size = 15

                leading = 21

            else:

                font_size = 13

                leading = 19

            elements.append(
                BidiTextBlock(
                    text=line,

                    font_name=font_name,

                    font_size=font_size,

                    leading=leading,

                    space_after=8
                )
            )

            continue

        # =================================================
        # Bullet
        # =================================================

        if is_bullet:

            elements.append(
                BidiListItem(
                    text=line,

                    marker="•",

                    font_name=font_name,

                    font_size=11,

                    leading=17,

                    space_after=5
                )
            )

            continue

        # =================================================
        # Numbered list
        # =================================================

        if number:

            elements.append(
                BidiListItem(
                    text=line,

                    marker=f"{number}.",

                    font_name=font_name,

                    font_size=11,

                    leading=17,

                    space_after=5
                )
            )

            continue

        # =================================================
        # Normal paragraph
        # =================================================

        elements.append(
            BidiTextBlock(
                text=line,

                font_name=font_name,

                font_size=11,

                leading=17,

                space_after=6
            )
        )

    # =====================================================
    # Build PDF
    # =====================================================

    document.build(
        elements
    )

    buffer.seek(0)

    return buffer


# =========================================================
# WORD RTL HELPERS
# =========================================================

def set_paragraph_rtl(paragraph):
    """
    Set Word paragraph direction to RTL.
    """

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    paragraph_properties = (
        paragraph._p.get_or_add_pPr()
    )

    existing_bidi = paragraph_properties.find(
        qn("w:bidi")
    )

    if existing_bidi is None:

        bidi = OxmlElement(
            "w:bidi"
        )

        bidi.set(
            qn("w:val"),
            "1"
        )

        paragraph_properties.append(
            bidi
        )


def set_run_rtl(run):
    """
    Set only an Arabic Word run to RTL.

    English runs inside the same paragraph stay LTR.
    """

    run_properties = (
        run._r.get_or_add_rPr()
    )

    existing_rtl = run_properties.find(
        qn("w:rtl")
    )

    if existing_rtl is None:

        rtl = OxmlElement(
            "w:rtl"
        )

        rtl.set(
            qn("w:val"),
            "1"
        )

        run_properties.append(
            rtl
        )

def set_run_ltr(run):
    """
    Explicitly force a Word run to LTR.
    """

    run_properties = run._r.get_or_add_rPr()

    rtl = run_properties.find(
        qn("w:rtl")
    )

    if rtl is None:
        rtl = OxmlElement("w:rtl")
        run_properties.append(rtl)

    rtl.set(
        qn("w:val"),
        "0"
    )

# =========================================================
# WORD MIXED ARABIC / ENGLISH
# =========================================================

def split_directional_runs(text: str):
    """
    Generic Unicode-based splitter for mixed Arabic/English text.

    Handles:
        - Arabic
        - English
        - Numbers
        - Parentheses
        - Brackets
        - Quotes
        - Colons
        - Slashes
        - Dashes
        - General punctuation

    No lecture-specific words are used.
    """

    if not text:
        return []

    base_direction = detect_base_direction(text)

    # -----------------------------------------------------
    # Strong direction of a character
    # -----------------------------------------------------

    def strong_direction(char):

        bidi = unicodedata.bidirectional(char)

        # Arabic / RTL scripts
        if bidi in ("R", "AL"):
            return "R"

        # Latin + numbers
        if bidi in ("L", "EN", "AN"):
            return "L"

        # Neutral character
        return None

    length = len(text)

    raw_directions = [
        strong_direction(char)
        for char in text
    ]

    # -----------------------------------------------------
    # Previous strong direction
    # -----------------------------------------------------

    previous_direction = [None] * length
    previous_distance = [None] * length

    last_direction = None
    last_index = None

    for i in range(length):

        previous_direction[i] = last_direction

        if last_index is not None:
            previous_distance[i] = i - last_index

        if raw_directions[i] is not None:
            last_direction = raw_directions[i]
            last_index = i

    # -----------------------------------------------------
    # Next strong direction
    # -----------------------------------------------------

    next_direction = [None] * length
    next_distance = [None] * length

    last_direction = None
    last_index = None

    for i in range(length - 1, -1, -1):

        next_direction[i] = last_direction

        if last_index is not None:
            next_distance[i] = last_index - i

        if raw_directions[i] is not None:
            last_direction = raw_directions[i]
            last_index = i

    # -----------------------------------------------------
    # Resolve punctuation / spaces / brackets
    # -----------------------------------------------------

    resolved = []

    for i, char in enumerate(text):

        direction = raw_directions[i]

        if direction is not None:
            resolved.append(direction)
            continue

        prev_dir = previous_direction[i]
        next_dir = next_direction[i]

        prev_dist = previous_distance[i]
        next_dist = next_distance[i]

        category = unicodedata.category(char)

        # -------------------------------------------------
        # Opening brackets / opening quotation marks
        #
        # (
        # [
        # {
        #
        # Attach them to what comes AFTER.
        #
        # Example:
        # عربي (English)
        #
        # "(" becomes LTR.
        # -------------------------------------------------

        if category in ("Ps", "Pi"):

            direction = (
                next_dir
                or prev_dir
                or base_direction
            )

        # -------------------------------------------------
        # Closing brackets / closing quotation marks
        #
        # )
        # ]
        # }
        #
        # Attach them to what came BEFORE.
        # -------------------------------------------------

        elif category in ("Pe", "Pf"):

            direction = (
                prev_dir
                or next_dir
                or base_direction
            )

        # -------------------------------------------------
        # Spaces
        # -------------------------------------------------

        elif char.isspace():

            # If both sides have the same direction,
            # keep the space with them.
            if (
                prev_dir
                and next_dir
                and prev_dir == next_dir
            ):
                direction = prev_dir

            # Otherwise prefer the previous text.
            else:
                direction = (
                    prev_dir
                    or next_dir
                    or base_direction
                )

        # -------------------------------------------------
        # Other punctuation:
        #
        # :
        # ;
        # ,
        # .
        # /
        # -
        # _
        # etc.
        # -------------------------------------------------

        else:

            if (
                prev_dir
                and next_dir
                and prev_dir == next_dir
            ):

                direction = prev_dir

            elif prev_dir and next_dir:

                # Attach neutral punctuation to
                # the nearest strong text.
                if (
                    prev_dist is not None
                    and next_dist is not None
                ):

                    if prev_dist <= next_dist:
                        direction = prev_dir
                    else:
                        direction = next_dir

                else:

                    direction = (
                        prev_dir
                        or next_dir
                    )

            else:

                direction = (
                    prev_dir
                    or next_dir
                    or base_direction
                )

        resolved.append(direction)

    # -----------------------------------------------------
    # Merge characters with the same direction into runs
    # -----------------------------------------------------

    parts = []

    current_text = text[0]
    current_direction = resolved[0]

    for i in range(1, length):

        direction = resolved[i]

        if direction == current_direction:

            current_text += text[i]

        else:

            parts.append(
                (
                    current_text,
                    current_direction
                )
            )

            current_text = text[i]
            current_direction = direction

    parts.append(
        (
            current_text,
            current_direction
        )
    )

    return parts


def add_directional_text(
    paragraph,
    text: str,
    bold: bool = False
):
    """
    Add mixed Arabic/English text using real Word
    RTL/LTR run properties.

    No hidden Unicode direction characters are inserted.
    """

    directional_parts = split_directional_runs(
        text
    )

    for segment, direction in directional_parts:

        run = paragraph.add_run(
            segment
        )

        run.bold = bold

        if direction == "R":

            set_run_rtl(
                run
            )

        else:

            set_run_ltr(
                run
            )

# =========================================================
# WORD MARKDOWN
# =========================================================

def add_markdown_runs(
    paragraph,
    text: str
):
    """
    Convert:

        **CMMI**

    into an actual bold Word run.

    Also preserves Arabic / English direction.
    """

    parts = re.split(
        r"(\*\*.*?\*\*)",
        text
    )

    for part in parts:

        if not part:

            continue

        # =================================================
        # Bold Markdown
        # =================================================

        if (
            part.startswith("**")
            and part.endswith("**")
        ):

            content = part[2:-2]

            add_directional_text(
                paragraph,
                content,
                bold=True
            )

        # =================================================
        # Normal
        # =================================================

        else:

            add_directional_text(
                paragraph,
                part,
                bold=False
            )


# =========================================================
# WORD EXPORT
# =========================================================

def generate_word(summary_text: str):
    """
    Generate DOCX from stored Summary.

    Supports:
        - Arabic RTL
        - English LTR
        - Mixed Arabic / English
        - Markdown headings
        - Markdown bold
        - Bullets
        - Numbered lists
    """

    document = Document()

    document.core_properties.title = (
        "Lecture Summary"
    )

    # =====================================================
    # Process Summary
    # =====================================================

    for raw_line in summary_text.splitlines():

        line = raw_line.strip()

        # =================================================
        # Empty line
        # =================================================

        if not line:

            document.add_paragraph()

            continue

        # =================================================
        # Heading level 3
        # =================================================

        if line.startswith("### "):

            text = line[4:]

            paragraph = document.add_heading(
                "",
                level=3
            )

            add_markdown_runs(
                paragraph,
                text
            )

            if detect_base_direction(text) == "R":

                set_paragraph_rtl(
                    paragraph
                )

            continue

        # =================================================
        # Heading level 2
        # =================================================

        elif line.startswith("## "):

            text = line[3:]

            paragraph = document.add_heading(
                "",
                level=2
            )

            add_markdown_runs(
                paragraph,
                text
            )

            if detect_base_direction(text) == "R":

                set_paragraph_rtl(
                    paragraph
                )

            continue

        # =================================================
        # Heading level 1
        # =================================================

        elif line.startswith("# "):

            text = line[2:]

            paragraph = document.add_heading(
                "",
                level=1
            )

            add_markdown_runs(
                paragraph,
                text
            )

            if detect_base_direction(text) == "R":

                set_paragraph_rtl(
                    paragraph
                )

            continue

        # =================================================
        # Bullet
        # =================================================

        elif line.startswith(
            (
                "- ",
                "* ",
                "• "
            )
        ):

            text = line[2:].strip()

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            add_markdown_runs(
                paragraph,
                text
            )

            if detect_base_direction(text) == "R":

                set_paragraph_rtl(
                    paragraph
                )

            continue

        # =================================================
        # Numbered list
        # =================================================

        numbered_match = re.match(r"^(\d+)\.\s+(.*)",line)


        if numbered_match:

            number = numbered_match.group(1)
            text = numbered_match.group(2)

            paragraph = document.add_paragraph()

            full_text = f"{number}. {text}"

            add_markdown_runs(
                paragraph,
                full_text
            )

            if detect_base_direction(text) == "R":
                set_paragraph_rtl(
                    paragraph
                )

            continue
        # =================================================
        # Normal paragraph
        # =================================================

        paragraph = document.add_paragraph()

        add_markdown_runs(
            paragraph,
            line
        )

        if detect_base_direction(line) == "R":

            set_paragraph_rtl(
                paragraph
            )

    # =====================================================
    # Save Word to memory
    # =====================================================

    buffer = io.BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    return buffer